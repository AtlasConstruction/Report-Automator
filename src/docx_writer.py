from docx import Document
from docx.shared import Inches

class DocxWriter:
    def __init__(self):
        self.document = Document()
    
    def create_table(self, data):
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        for item in data:
            row_cells = table.add_row().cells
            # Add image to the first column
            row_cells[0].paragraphs[0].add_run().add_picture(
                item['path'], 
                width=Inches(2)
            )
            # Add text to the second column
            row_cells[1].text = item['text']
    
    def save(self, file_path):
        self.document.save(file_path)